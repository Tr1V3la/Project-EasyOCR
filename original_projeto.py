#bibliotecas necessarias ou que foram utilizadas em algum momento
import tkinter as tk
from tkinter import ttk, filedialog
from PIL import ImageTk, Image
import PyPDF2
import docx
import easyocr
import pyttsx3
import cv2
from googletrans import Translator
from tkinter import Toplevel
from langdetect import detect
import numpy as np
from io import BytesIO
import requests
import customtkinter as ctk
from gtts import gTTS
import tempfile
import os
os.environ['PYGAME_HIDE_SUPPORT_PROMPT'] = '1'
import pygame
import time
from pydub import AudioSegment
from pydub.playback import play

IDIOMAS_VALIDOS = ['en', 'pt']

#Garantindo que não tenha nada na variavel
janela_previa = None
txt_previa = None
txt_previa_janela_previa = None
scrollbar = None
bnt_extract = None
arquivo_selecionado = False
imagem_selecionada = None
caminho_arquivo = None
cap = None
combo_idioma = None
lbl_selecao = None
speaker = None
btn_capturar = None
frame_atual = None
imagem_capturada = None
camera_aberta = False
tooltip_window = None
is_playing = False
is_paused = False
temp_file = None




url = "https://wallpapercave.com/wp/wp4363675.jpg"  #URL da imagem

#Solicitação GET para baixar a imagem
response = requests.get(url)

# Verifique se a solicitação foi bem-sucedida
if response.status_code == 200:
    # Armazene o conteúdo binário da imagem na variável img_fundo_aplicar
    img_fundo_aplicar = response.content
else:
    print("Falha ao baixar a imagem")

image_data = BytesIO(img_fundo_aplicar)


language_mapping = {
    "Inglês": "en",
    "Português": "pt"
}
#Campo para selecionar o idioma
def selecionar_idioma(event):
    global idioma_selecionado, combo_idioma, bnt_extract, bnt_play_pause, bnt_velocidade
    idioma_selecionado = language_mapping[combo_idioma.get()]
    combo_idioma.selection_clear()
    lbl_selecao.config(text="Idioma selecionado: " + idioma_selecionado)
    
    if idioma_selecionado in ["en", "pt"]:
        # Habilitar botões
        if bnt_extract is not None:
            bnt_extract["state"] = "normal"
        if bnt_play_pause is not None:
            bnt_play_pause["state"] = "normal"
        if bnt_velocidade is not None:
            bnt_velocidade["state"] = "normal"
        
        # Extrair o texto do arquivo ou imagem
        if caminho_arquivo is not None:
            texto_traduzido, _ = extract_text_from_file(caminho_arquivo, txt_previa_janela_previa)
        elif imagem_selecionada is not None:
            texto_traduzido, _ = extract_text_from_image(imagem_selecionada, idioma_selecionado)
        elif imagem_capturada is not None:  # Verifica se uma imagem foi capturada pela webcam
            texto_traduzido, _ = extract_text_from_image(imagem_capturada, idioma_selecionado)
        else:
            print("No file or image selected.")
            return
        
    else:
        print("Por favor, selecione um idioma.")
        # Desabilitar botões
        if bnt_extract is not None:
            bnt_extract["state"] = "disabled"
        if bnt_play_pause is not None:
            bnt_play_pause["state"] = "disabled"
        if bnt_velocidade is not None:
            bnt_velocidade["state"] = "disabled"



# abrir a janela de seleção de arquivo
def selecionar_arquivo():
    global arquivo_selecionado, caminho_arquivo, btn_selecionar_arquivo
    if arquivo_selecionado:
        # Se um arquivo já foi selecionado, então "Retirar Arquivo" foi clicado
        arquivo_selecionado = False
        caminho_arquivo = None
        btn_selecionar_arquivo.configure(text="Selecionar Arquivo")
        txt_previa.delete(1.0, tk.END)  # Limpa a janela de visualização
        btn_aplicar.configure(state=tk.DISABLED)  # Desativa o botão 'Aplicar'
        btn_selecionar_imagem.configure(state=tk.NORMAL)  # Ativa o botão 'Selecionar Imagem'
        btn_capturar_imagem.configure(state=tk.NORMAL)  # Ativa o botão 'Câmera'
    else:
        # Se nenhum arquivo foi selecionado, então "Selecionar Arquivo" foi clicado
        arquivo = filedialog.askopenfilename(
            filetypes=(("PDF Files", "*.pdf"), ("Word Documents", "*.docx"))
        )
        if arquivo:
            print("Arquivo selecionado:", arquivo)
            exibir_previa_arquivo(arquivo)
            arquivo_selecionado = True
            caminho_arquivo = arquivo  # Armazenar o caminho do arquivo
            btn_selecionar_arquivo.configure(text="Retirar Arquivo")
            btn_aplicar.configure(state=tk.NORMAL)  # Ativa o botão 'Aplicar' 
            btn_selecionar_imagem.configure(state=tk.DISABLED)  # Desativa o botão 'Selecionar Imagem'
            btn_capturar_imagem.configure(state=tk.DISABLED)  # Desativa o botão 'Câmera'


# exibindo uma previa do arquivo para pdf limitado na primeira pagina e para word limitado em paragrafos
def exibir_previa_arquivo(arquivo):
    global txt_previa, txt_previa_janela_previa
    page_text = ""  # Initialize page_text to an empty string
    
    if janela_previa:
        txt_previa_janela_previa.delete(1.0, tk.END)
        txt_previa_janela_previa.insert(
            tk.END, f"Visualização Prévia do Arquivo: {arquivo}\n\n")
    else:
        txt_previa.delete(1.0, tk.END)
        txt_previa.insert(
            tk.END, f"Visualização Prévia do Arquivo: {arquivo}\n\n")

    if arquivo.lower().endswith(".pdf"):
        pdf_file = open(arquivo, "rb")
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        first_page = pdf_reader.pages[0]
        page_text = first_page.extract_text()
        if janela_previa:
            txt_previa_janela_previa.insert(tk.END, page_text)
        else:
            txt_previa.insert(tk.END, page_text)

    elif arquivo.lower().endswith(".docx"):
        doc = docx.Document(arquivo)
        paragraphs = doc.paragraphs[:10]
        for paragraph in paragraphs:
            if janela_previa:
                txt_previa_janela_previa.insert(tk.END, paragraph.text + "\n")
            else:
                txt_previa.insert(tk.END, paragraph.text + "\n")
        page_text = "".join([paragraph.text for paragraph in paragraphs])

    # Detect the language of the text only if page_text has enough content
    cleaned_text = page_text.strip()
    if len(cleaned_text) > 10:
        try:
            language = detect(page_text)
            print("Detected language:", language)
        except Exception as e:
            print(f"Error in language detection: {e}")
    else:
        print("Text is too short for language detection.")
        


# exibindo uma previa da imagem selecionada
def exibir_previa_imagem(input_image):
    global txt_previa, txt_previa_janela_previa, img_label
    
    # Verificar se a entrada é um caminho de arquivo ou um numpy.ndarray
    if isinstance(input_image, str):
        # Se for um caminho de arquivo, abra-o
        img = Image.open(input_image)
    elif isinstance(input_image, np.ndarray):
        # Se for um array numpy, converta-o em imagem PIL
        img = Image.fromarray(cv2.cvtColor(input_image, cv2.COLOR_BGR2RGB))
    else:
        raise ValueError("A entrada deve ser um caminho de arquivo ou um numpy.ndarray!")

    # Redimensionar a imagem
    img = img.resize((300, 200))
    img_preview = ImageTk.PhotoImage(img)
    img_label.configure(image=img_preview)
    img_label.image = img_preview

    if janela_previa:
        txt_previa_janela_previa.delete(1.0, tk.END)
        txt_previa_janela_previa.insert(
            tk.END, f"Visualização Prévia da Imagem: {input_image}\n\n")
    else:
        txt_previa.delete(1.0, tk.END)
        txt_previa.insert(
            tk.END, f"Visualização Prévia da Imagem: {input_image}\n\n")


# Função para abrir a webcam
def abrir_webcam():
    global frame_atual, imagem_capturada, image_data, captura_ativa, lbl_camera, vc, camera_aberta, btn_capturar_imagem, btn_tirar_foto

    frame_camera = None

    if camera_aberta:
        captura_ativa = False  # Desativa o loop de atualização da imagem
        if vc:
            vc.release()  # Libera os recursos da webcam
            vc = None  # Remove qualquer referência restante
        if lbl_camera:
            lbl_camera.pack_forget()  # Remove a janela de exibição da webcam
        camera_aberta = False
        btn_capturar_imagem["text"] = "Camera"
        btn_tirar_foto.configure(state=tk.DISABLED)
        btn_tirar_foto["text"] = "Tirar foto"
        imagem_capturada = None
        btn_selecionar_arquivo.configure(state=tk.NORMAL)
        btn_selecionar_imagem.configure(state=tk.NORMAL)
        
        # Limpar a imagem de prévia
        img_label.configure(image="")
             
        # Limpar o widget txt_previa
        txt_previa.delete(1.0, tk.END)
    else:
        captura_ativa = True
        frame_camera = tk.Frame(janela)
        frame_camera.place(relx=0.8, rely=0.73, anchor='center', width=350, height=300)

        lbl_camera = tk.Label(frame_camera)
        lbl_camera.pack(fill="both", expand="yes")

        vc = cv2.VideoCapture(0)
        frame_atual = None
        atualizar_frame()
        btn_tirar_foto.configure(state=tk.NORMAL)

        camera_aberta = True
        btn_capturar_imagem["text"] = "Fechar camera"
        btn_selecionar_arquivo.configure(state=tk.DISABLED)  # Desativa o botão 'Selecionar Arquivo'
        btn_selecionar_imagem.configure(state=tk.DISABLED)  # Desativa o botão 'Selecionar Imagem'

    return frame_camera


def atualizar_frame():
    global frame_atual, imagem_capturada, captura_ativa, lbl_camera, vc
    if captura_ativa:
        rval, frame = vc.read()
        if rval:
            # Inverte a imagem  (efeito de espelho)
            #frame = cv2.flip(frame, 1)
            # Converta a imagem da câmera para um formato que pode ser usado no Tkinter
            frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
            img = Image.fromarray(frame)
            imgtk = ImageTk.PhotoImage(image=img)
            lbl_camera.imgtk = imgtk
            lbl_camera.configure(image=imgtk)
            if imagem_capturada is None:  # Só atualiza frame_atual se imagem_capturada estiver vazia
                frame_atual = frame 
        # Atualize o frame a cada 10 ms
        lbl_camera.after(10, atualizar_frame) 
        
def capturar_imagem():
    global frame_atual, imagem_capturada, captura_ativa
    if captura_ativa:
        # A captura está ativa, então tiramos a foto
        if frame_atual is not None:
            # Aumenta a nitidez da imagem
            kernel = np.array([[0, -1, 0],
                               [-1, 5,-1],
                               [0, -1, 0]])
            imagem_capturada = cv2.filter2D(frame_atual, -1, kernel)

            # Ajusta o brilho e o contraste da imagem capturada
            alpha = 1.5  # ajuste de contraste
            beta = 30    # ajuste de brilho
            imagem_capturada = cv2.convertScaleAbs(imagem_capturada, alpha=alpha, beta=beta)

            # Agora imagem_capturada é um array NumPy que você pode passar para outras funções
            exibir_previa_imagem(imagem_capturada)

            btn_aplicar.configure(state=tk.NORMAL)
            btn_tirar_foto.configure(text="Retirar foto")
            captura_ativa = False
    else:
        # A captura está inativa, então reativamos a captura
        imagem_capturada = None
        btn_aplicar.configure(state=tk.DISABLED)
        btn_tirar_foto.configure(text="Tirar foto")
        captura_ativa = True
        atualizar_frame()



def translate_text(text, target_language):
    if text is None or text.strip() == "":
        print("Nenhum texto para traduzir.")
        return None
    if target_language not in ['en', 'pt']:
        raise ValueError(f'Invalid language code: {target_language}')
    
    translator = Translator()
    try:
        translation = translator.translate(text, dest=target_language)
        translated_text = translation.text
        if not translated_text or translated_text.isspace():
            print("Translation returned empty or whitespace-only string.")
            return None
        return translated_text
    except Exception as e:
        print(f"Error while translating: {e}")
        return None


# Criando botão para voltar a pagina incial
def voltar_para_home():
    global janela_previa, arquivo_selecionado, img_fundo, temp_file, caminho_arquivo, imagem_capturada, camera_aberta

    # Pare a música e exclua o arquivo temporário, se necessário
    if temp_file and pygame.mixer.get_init() and pygame.mixer.music.get_busy():
        pygame.mixer.music.stop()
        pygame.mixer.quit()  # Finalize o mixer
        for _ in range(5):  # Tente excluir até 5 vezes
            try:
                os.unlink(temp_file.name)
                break
            except Exception as e:
                print(f'Erro ao excluir arquivo temporário (tentativa {_ + 1}): {e}')
                time.sleep(0.5)  # Espere por meio segundo antes de tentar novamente

    # Limpar a janela de visualização antes de destruir a janela
    if janela_previa:
        txt_previa_janela_previa.delete(1.0, tk.END)
        img_label_janela_previa.configure(image="")
        janela_previa.destroy()  # Destrói a janela de pré-visualização

    janela.deiconify()
    janela_previa = None
    arquivo_selecionado = False
    btn_aplicar.configure(state=tk.DISABLED)

    # Limpar a imagem capturada da webcam, se houver
    imagem_capturada = None

    # Se a webcam estiver aberta, feche-a
    if camera_aberta:
        abrir_webcam()

    # Limpar a janela principal, se necessário
    txt_previa.delete(1.0, tk.END)
    img_label.configure(image="")
    caminho_arquivo = None
    # Redefinir os estados dos botões
    btn_selecionar_arquivo.configure(state=tk.NORMAL, text="Selecionar Arquivo")
    btn_selecionar_imagem.configure(state=tk.NORMAL, text="Selecionar Imagem")
    btn_capturar_imagem.configure(state=tk.NORMAL, text="Câmera")



def check_music_status_adjusted():
    global temp_file, is_paused
    if not pygame.mixer.music.get_busy() and not is_paused and temp_file:
        os.remove(temp_file.name)
        temp_file = None
    else:
        # Se a música estiver pausada ou ainda estiver tocando, verifique novamente em 1 segundo
        janela.after(1000, check_music_status_adjusted)
        
def start_playback():
    global is_playing, temp_file
    if not is_playing and temp_file:  # Se houver um arquivo temporário disponível e o áudio não estiver tocando
        pygame.mixer.music.load(temp_file.name)
        pygame.mixer.music.play()
        bnt_play_pause["text"] = " ▶┃"  # Atualiza o texto do botão para mostrar o ícone de play
        is_playing = True
        # Chamando a função check_music_status_adjusted após ter certeza de que a música começou a ser reproduzida
        janela.after(1000, check_music_status_adjusted)
        
def reproduzir_pausar_fala():
    global is_playing, is_paused
    
    # Se o áudio estiver tocando e não estiver pausado
    if is_playing and not is_paused:
        pygame.mixer.music.pause()
        bnt_play_pause["text"] = " ⏸"  # Atualiza o texto do botão para mostrar o ícone de pause
        is_paused = True 
        return

    # Se o áudio estiver pausado
    if is_playing and is_paused:
        pygame.mixer.music.unpause()
        bnt_play_pause["text"] = " ▶┃"  # Atualiza o texto do botão para mostrar o ícone de play
        is_paused = False 
        return

    # Se o áudio não estiver tocando e não estiver pausado (caso inicial)
    if not is_playing and not is_paused:
        start_playback()



def extract_text_from_image(input_image, target_language):
    global txt_previa_janela_previa

    # Verificar se a entrada é um caminho de arquivo ou um numpy.ndarray
    if isinstance(input_image, str):
        print(f"Extraindo texto da imagem: \n {input_image}")  # Debug print
        # Carrega imagem em escala de cinza
        img = cv2.imread(input_image, 0)
    elif isinstance(input_image, np.ndarray):
        img = input_image
    else:
        raise ValueError("A entrada deve ser um caminho de arquivo ou um numpy.ndarray!")

    # Redimensiona imagem
    img = cv2.resize(img, None, fx=3, fy=3, interpolation=cv2.INTER_CUBIC)

    # Convertendo a imagem CV2 para PIL para ser usada no label tkinter
    img_for_display = Image.fromarray(img)
    img_preview = ImageTk.PhotoImage(img_for_display)
    img_label.configure(image=img_preview)
    img_label.image = img_preview

    supported_languages = ['pt', 'en']  # Adicione mais idiomas se necessário
    reader = easyocr.Reader(supported_languages)
    try:
        result = reader.readtext(img)
    except Exception as e:
        error_message = f"Erro ao ler texto da imagem: {e}"
        print(error_message)
        txt_previa_janela_previa.insert(tk.END, error_message) 
        return "", 0 

    # Obtendo o texto extraído e as confianças
    txt = ""
    confs = []
    for r in result:
        txt += r[1] + " "
        confs.append(r[2])

    # Verificando se nenhum texto foi reconhecido
    if not txt.strip():
        txt_previa_janela_previa.insert(tk.END, "Não foi reconhecido nenhum texto. Por favor, retire a foto novamente.")
        return "", 0

    # Traduzir o texto para o idioma selecionado
    translated_text = translate_text(txt, target_language)

    txt_previa_janela_previa.insert(tk.END,f"Texto traduzido: \n\n {translated_text}") 

    # Calculando e exibindo a confiança média
    confianca_media = sum(confs) / len(confs) if confs else 0
    print(f"Confiança média: {confianca_media}")

    return translated_text, confianca_media



def aplicar_selecao():
    global arquivo_selecionado, imagem_selecionada, caminho_arquivo, imagem_capturada
    
    # Limpar a janela de texto da janela de prévia (se aplicável)
    if janela_previa:
        txt_previa_janela_previa.delete(1.0, tk.END)
    else:
        txt_previa.delete(1.0, tk.END)

    if caminho_arquivo is not None:
        abrir_janela_previa()
        extract_text_from_file(caminho_arquivo, txt_previa_janela_previa)
        arquivo_selecionado = False
    elif imagem_selecionada:
        abrir_janela_previa()
        target_language = combo_idioma.get()
        if target_language not in ["Inglês", "Português"]:
            print("Invalid language selection.")
            return
        elif target_language == "Inglês":
            target_language = "en"
        elif target_language == "Português":
            target_language = "pt"
        translated_text, _ = extract_text_from_image(imagem_selecionada, target_language)
        txt_previa_janela_previa.insert(tk.END, translated_text)
        imagem_selecionada = None
    elif imagem_capturada is not None:
        abrir_janela_previa()
        target_language = combo_idioma.get()
        
        # Verificação e seleção do idioma
        if target_language not in ["Inglês", "Português"]:
            print("Invalid language selection.")
            return
        elif target_language == "Inglês":
            target_language = "en"
        elif target_language == "Português":
            target_language = "pt"
        
        # Extração de texto da imagem usando EasyOCR
        translated_text, _ = extract_text_from_image(imagem_capturada, target_language)
        txt_previa_janela_previa.insert(tk.END, translated_text)
        
        # Limpar a imagem capturada e fechar a webcam, se necessário
        imagem_capturada = None
        if camera_aberta:
            abrir_webcam()
        
        # Desativar o botão "Aplicar" após o processamento
        btn_aplicar.configure(state=tk.DISABLED)

def extract_text_from_file(caminho_arquivo, text_widget, preview_mode=False):
    global txt_previa_janela_previa
    # Limpar o widget de texto
    txt_previa_janela_previa.delete(1.0, tk.END)
    # Inicializar uma variável para armazenar todo o texto extraído
    all_text = ""
    # Exibir o conteúdo do arquivo
    if caminho_arquivo is not None:
        if caminho_arquivo.lower().endswith(".pdf"):
            pdf_file = open(caminho_arquivo, "rb")
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                txt_previa_janela_previa.insert(tk.END, page_text)
                all_text += page_text  # Adicionar texto à variável all_text
        elif caminho_arquivo.lower().endswith(".docx"):
            doc = docx.Document(caminho_arquivo)
            for paragraph in doc.paragraphs:
                txt_previa_janela_previa.insert(tk.END, paragraph.text + "\n")
                all_text += paragraph.text + "\n"  # Adicionar texto à variável all_text
    # Retornar todo o texto extraído
    return all_text, None



def translate_file(file_path, target_language):
    text, _ = extract_text_from_file(file_path)
    if text is None or text.strip() == "":
        print("Nenhum texto extraído para traduzir.")
        return None
    translated_text = translate_text(text, target_language)
    return translated_text

    
# Exibindo a imagem selcionado
def selecionar_imagem():
    global arquivo_selecionado, imagem_selecionada, btn_selecionar_imagem
    if imagem_selecionada:
        # Se uma imagem já foi selecionada, então "Retirar Imagem" foi clicado
        imagem_selecionada = None
        btn_selecionar_imagem.configure(text="Selecionar Imagem")
        txt_previa.delete(1.0, tk.END)  # Limpa a janela de visualização
        img_label.configure(image="")  # Limpa a label da imagem
        btn_aplicar.configure(state=tk.DISABLED)  # Desativa o botão 'Aplicar'
        btn_selecionar_arquivo.configure(state=tk.NORMAL)  # Ativa o botão 'Selecionar Arquivo'
        btn_capturar_imagem.configure(state=tk.NORMAL)  # Ativa o botão 'Câmera'
    else:
        # Se nenhuma imagem foi selecionada, então "Selecionar Imagem" foi clicado
        imagem = filedialog.askopenfilename(
            filetypes=(("JPEG Files", "*.jpg"), ("PNG Files", "*.png"))
        )
        if imagem:
            print("Imagem selecionada:", imagem)
            imagem_selecionada = imagem
            exibir_previa_imagem(imagem)
            arquivo_selecionado = True
            btn_selecionar_imagem.configure(text="Retirar Imagem")
            btn_aplicar.configure(state=tk.NORMAL)
            btn_selecionar_arquivo.configure(state=tk.DISABLED)  # Desativa o botão 'Selecionar Arquivo'
            btn_capturar_imagem.configure(state=tk.DISABLED)  # Desativa o botão 'Câmera'

def alternar_texto():
    global estado_atual 
    estado_atual = (estado_atual + 1) % len(estados_velocidade)
    bnt_velocidade["text"] = estados_velocidade[estado_atual]

    pygame.mixer.music.stop()
    extrair_audio()

def speak_text(texto_traduzido):
    global stop_speaking, estado_atual
    engine = pyttsx3.init()

    # Ajustando a velocidade da fala com base no estado atual
    speed_multiplier = 100
    speed_addition = 50 if estado_atual == 0 else 0  # Aumenta a velocidade em 50 quando estado_atual é 0
    speed = speed_multiplier * (estado_atual + 1) + speed_addition  # Velocidade da fala (150, 200, 300) quando estados são ("1x", "2x", "3x")
    engine.setProperty('rate', speed)

    sentences = texto_traduzido.split('. ')
    for sentence in sentences:
        if stop_speaking:
            break
        engine.say(sentence)
        engine.runAndWait()

# alterar o status do botão de velocidade
estados_velocidade = ["1x", "2x", "3x"]
estado_atual = 0

temp_file = None

def extrair_audio():
    global caminho_arquivo, imagem_selecionada, txt_previa_janela_previa, idioma_selecionado, estado_atual, temp_file, is_playing

    # Verificar se o idioma selecionado é válido
    if idioma_selecionado not in IDIOMAS_VALIDOS:
        idioma_selecionado = 'en'  # Opção 1: Definir um valor padrão

    # Limpar o widget txt_previa_janela_previa
    txt_previa_janela_previa.delete(1.0, tk.END)

    if caminho_arquivo is not None:
        # Extrai o texto do arquivo
        texto_extraido, _ = extract_text_from_file(caminho_arquivo, txt_previa_janela_previa)

    elif imagem_selecionada is not None:
        # Extrai o texto da imagem
        texto_extraido, _ = extract_text_from_image(imagem_selecionada, idioma_selecionado)

    else:
        print("No file or image selected.")
        return

    # Traduz o texto extraído para o idioma selecionado
    texto_traduzido = translate_text(texto_extraido, idioma_selecionado)

    # Limpar janela
    txt_previa_janela_previa.delete(1.0, tk.END)

    # Insere o texto traduzido na janela prévia
    if texto_traduzido is not None and texto_traduzido.strip() != "":
        txt_previa_janela_previa.insert(tk.END, "Texto Traduzido:\n\n")
        txt_previa_janela_previa.insert(tk.END, texto_traduzido)
    else:
        print("Nenhum texto para traduzir.")
        return

    # Converte o texto em áudio usando gTTS e salva em um arquivo temporário
    if texto_traduzido is not None and texto_traduzido.strip() != "":
        tts = gTTS(text=texto_traduzido, lang=idioma_selecionado)
        if temp_file:  # Se já existir um arquivo temporário anterior, exclua-o
            pygame.mixer.quit()  # Finalize o mixer
            if os.path.exists(temp_file.name):
                os.unlink(temp_file.name)

        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".mp3")
        tts.save(temp_file.name)
        temp_file.close()
    else:
        print("Nenhum texto para converter em áudio.")
        return
        
    audio = AudioSegment.from_mp3(temp_file.name)
    playback_speeds = [1.0, 1.4, 1.7] #definição das velocidades
    speed = playback_speeds[estado_atual]
    
    if speed != 1.0:  # Se a velocidade for diferente de 1x, ajuste a velocidade
        audio = audio.speedup(playback_speed=speed)

    # Salvar o áudio ajustado em um novo arquivo temporário
    temp_file_adjusted = tempfile.NamedTemporaryFile(delete=False, suffix=".mp3")
    audio.export(temp_file_adjusted.name, format="mp3")
    temp_file_adjusted.close()

    # Iniciar a reprodução usando pygame
    pygame.mixer.init()
    pygame.mixer.music.load(temp_file_adjusted.name)
    pygame.mixer.music.play()
    
    is_playing = True



# Atualizar a chamada para usar a nova função
def alternar_texto():
    global estado_atual 
    estado_atual = (estado_atual + 1) % len(estados_velocidade)
    bnt_velocidade["text"] = estados_velocidade[estado_atual]

    pygame.mixer.music.stop()
    extrair_audio()





def safe_remove(file_name, max_attempts=10, delay=0.5):
    """Tenta excluir um arquivo várias vezes, esperando entre as tentativas."""
    for _ in range(max_attempts):
        try:
            os.remove(file_name)
            return  # O arquivo foi excluído com sucesso
        except PermissionError:
            time.sleep(delay)  # Espera antes de tentar novamente


def on_enter_previa(event):
    global tooltip_window
    x = janela_previa.winfo_x() + button_previa.winfo_x() + button_previa.winfo_width() + 5
    y = janela_previa.winfo_y() + button_previa.winfo_y() + button_previa.winfo_height()
    tooltip_window = tk.Toplevel(janela_previa)
    tooltip_window.wm_overrideredirect(True)  # Remove window decorations
    tooltip_window.wm_geometry(f"+{x}+{y}")
    texto_tooltip_previa = ("Selecionando o idioma é liberado os botões de controle.\n"
                            "O botão extrair para audio serve tano para imagem quanto para o texto \n"
                            "Para os casos de imagem, o conteudo será mostrado na janela após o programa extrair o texto da imagem\n"
                            "O botão de pause/play, funciona para conseguir controlar o audio\n"
                            "Caso tenha concluido o uso e quer retornar o programa, pode apertar no botão Home\n"
                            "Ou pode solicitar para fechar a janela mesmo.\n"
                            "Utilize da melhor maneira que quiser")
    texto = tk.Text(tooltip_window, wrap=tk.WORD, bg="lightgray", fg="black", height=20, width=30)
    texto.insert(tk.END, texto_tooltip_previa)
    texto.config(state=tk.DISABLED)  # Disable editing
    texto.pack()



# janela quando apertar "Aplicar"
def abrir_janela_previa():
    global janela_previa, button_previa, txt_previa_janela_previa, bnt_play_pause, confianca_label, img_label_janela_previa, btn_aplicar_janela_previa, img_fundo_previa, bnt_velocidade, caminho_arquivo, combo_idioma, lbl_selecao
    janela.withdraw()  # Oculta a janela principal

    janela_previa = Toplevel(janela)
    janela_previa.title("Extração para audio")
    janela_previa.geometry("1000x600")
    janela_previa.resizable(False, False)

    # Carregar imagem de fundo da prévia
    imagem_fundo_previa = Image.open(image_data,)
    # dimensões da janela de prévia
    largura_janela_previa, altura_janela_previa = 1000, 600
    imagem_fundo_previa = imagem_fundo_previa.resize(
        (largura_janela_previa, altura_janela_previa), Image.LANCZOS)
    img_fundo_previa = ImageTk.PhotoImage(imagem_fundo_previa)

    lbl_fundo_previa = tk.Label(janela_previa, image=img_fundo)
    lbl_fundo_previa.place(x=0, y=0, relwidth=1, relheight=1)
    
    # Botão de interrogação na janela_previa
    button_previa = ctk.CTkButton(janela_previa, text="?", font=("Arial", 20), width=3, bg_color="#000000")
    button_previa.place(relx=1, y=0, anchor=tk.NE)  # Posiciona no canto superior direito
    button_previa.bind("<Enter>", on_enter_previa)
    button_previa.bind("<Leave>", on_leave)


    #Campo "selecionar idioma"
    idiomas = ["Inglês", "Português"]
    idioma_selecionado = None

    lbl_selecao = tk.Label(
        janela_previa, text="Idioma selecionado:", bg="#FFFFFF", fg="#000000", font=("Arial", 10)
    )
    lbl_selecao.pack(padx=10, pady=10, anchor="w")

    combo_idioma = ttk.Combobox(
        janela_previa, values=idiomas, state="readonly")
    combo_idioma.set("Selecione o idioma para o áudio")
    combo_idioma.pack(padx=10, pady=5, anchor="w")
    combo_idioma.configure(width=30)
    combo_idioma.bind("<<ComboboxSelected>>", selecionar_idioma)

    txt_previa_janela_previa = tk.Text(
        janela_previa,
        width=80,
        height=10,
        wrap=tk.WORD,
        bg="#282828",
        fg="#FFFFFF",
        font=("Arial", 10),
    )
    txt_previa_janela_previa.pack(
        padx=10, pady=10, anchor="w", fill=tk.BOTH, expand=True)

    # Exibir o conteúdo do arquivo
    extract_text_from_file(caminho_arquivo, txt_previa_janela_previa, preview_mode=False)


    scrollbar_janela_previa = tk.Scrollbar(txt_previa_janela_previa)
    scrollbar_janela_previa.pack(side=tk.RIGHT, fill=tk.Y)
    txt_previa_janela_previa.config(yscrollcommand=scrollbar_janela_previa.set)
    scrollbar_janela_previa.config(command=txt_previa_janela_previa.yview)

    img_label_janela_previa = tk.Label(janela_previa)
    img_label_janela_previa.pack(padx=10, pady=10, anchor="w")

    global bnt_extract
    # botão para extrair em audio
    bnt_extract = tk.Button(
        janela_previa,
        text="Extrair para Áudio",
        command= extrair_audio,
        width=20,
        height=2,
        bg="#FFFFFF",
        fg="#282828",
        state="disabled"  # Desabilita inicialmente
    )
    bnt_extract.pack(side="top", padx=10, pady=10, anchor="w")

    # botão de play ou pause
    bnt_play_pause = tk.Button(
        janela_previa,
        text=" ▶┃",
        command=reproduzir_pausar_fala,
        width=4,
        height=2,
        bg="#FFFFFF",
        fg="#282828",
        state=tk.DISABLED,
    )
    bnt_play_pause.pack(side="left", padx=10, pady=10, anchor="w")

    # botão de controle de velocidade
    bnt_velocidade = tk.Button(
        janela_previa,
        text=estados_velocidade[estado_atual],
        command=alternar_texto,
        width=8,
        height=2,
        bg="#FFFFFF",
        fg="#282828",
        state=tk.DISABLED,
    )
    bnt_velocidade.pack(side="left", padx=10, pady=10, anchor="w")

    # botão home
    btn_aplicar_janela_previa = tk.Button(
        janela_previa,
        text="Home",
        command=voltar_para_home,
        width=20,
        height=2,
        bg="#FFFFFF",
        fg="#282828",
    )
    btn_aplicar_janela_previa.pack(side="bottom", padx=10, pady=10, anchor="w")

    janela_previa.protocol("WM_DELETE_WINDOW", voltar_para_home)
    
        
    #Colocando o texto na tela Secundaria 
    coopyright_label = tk.Label(janela_previa, text="Copyright - Vinicius Cosmos  © 2023", font=("Arial", 10), fg="white", bg="#000000")
    coopyright_label.place(x=760, y=altura_janela - 20)

def on_enter(event):
    global tooltip_window
    x = janela.winfo_x() + button.winfo_x() + button.winfo_width() + 5
    y = janela.winfo_y() + button.winfo_y() + button.winfo_height()
    tooltip_window = tk.Toplevel(janela)
    tooltip_window.wm_overrideredirect(True) # Remove window decorations
    tooltip_window.wm_geometry(f"+{x}+{y}")
    texto_tooltip = ("Este App passa todo o conteudo inserido para aúdio\n"
                     "Para arquivos está limitado a PDF/Word\n"
                     "Para imagens está limitado a JPG/PNG\n"
                     "Webcam apenas irá funcionar caso contenha algum conteudo em texto para ser extraido.\n"
                     "Casos contrarios ao indicados o programa irá retornar uma mensagem de erro e zera todo o processo.\n"
                     "Clique nos botões para a função que deseja.")
    texto = tk.Text(tooltip_window, wrap=tk.WORD, bg="lightgray", fg="black", height=16, width=30)
    texto.insert(tk.END, texto_tooltip)
    texto.config(state=tk.DISABLED)  # Desativar edição
    texto.pack()
    
def on_leave(event):
    global tooltip_window
    tooltip_window.destroy()

def creditos_finais():
    # Esconder a janela principal
    janela.withdraw()

    # Criar uma nova janela
    janela_creditos = tk.Toplevel(janela)
    janela_creditos.title("Créditos")
    janela_creditos.geometry("1000x600")
    janela_creditos.resizable(False, False)
    janela_creditos.configure(bg="black") # Fundo preto

    # Texto centralizado
    texto_creditos = ("CRÉDITOS FINAIS\n\n\n""Pontifícia Universidade Católica de Campinas\n\n""Criado por: Vinícius Cosmos \n\n""Curso: Engenharia Elétrica\n\n\n\n""2023")
    lbl_creditos = tk.Label(janela_creditos, text=texto_creditos, font=("Arial", 20), fg="white", bg="black", wraplength=800)
    lbl_creditos.place(x=500, y=300, anchor=tk.CENTER) # Posiciona o texto no centro da janela
  
    # Função para fechar a janela de créditos e retornar à janela principal
    def voltar_para_home():
        janela_creditos.destroy()
        janela.deiconify() # Restaurar a janela principal

    # Botão "Home" para fechar a janela de créditos
    btn_home = tk.Button(
        janela_creditos,
        text="Home",
        command=voltar_para_home,
        width=20,
        height=2,
        bg="#FFFFFF",
        fg="#000000",
        relief=tk.RAISED,
        bd=5,
    )
    btn_home.place(x=10, y=550)


# Interface principal
janela = tk.Tk()
janela.title("Home")
janela.geometry("1000x600")
janela.resizable(False, False)

# Carregar imagem de fundo
imagem_fundo = Image.open(image_data)

# Especifique as dimensões da janela principal
largura_janela, altura_janela = 1000, 600
imagem_fundo = imagem_fundo.resize(
    (largura_janela, altura_janela), Image.LANCZOS)
img_fundo = ImageTk.PhotoImage(imagem_fundo)

lbl_fundo = tk.Label(janela, image=img_fundo)
lbl_fundo.place(x=0, y=0, relwidth=1, relheight=1)



# Botão de interrogação
button = ctk.CTkButton(janela, text="?", font=("Arial", 20), width=3, state=tk.DISABLED, bg_color="#000000")
button.place(relx=1, y=0, anchor=tk.NE) # Posiciona no canto superior direito
button.bind("<Enter>", on_enter)
button.bind("<Leave>", on_leave)


def redimensionar_txt_previa(event):
    txt_previa.config(width=60, height=10)


janela.bind("<Configure>", redimensionar_txt_previa)

txt_previa = tk.Text(
    janela,
    wrap=tk.WORD,
    bg="#282828",
    fg="#FFFFFF",
    font=("Arial", 10),
)
txt_previa.pack(padx=10, pady=10, anchor="w")

scrollbar = tk.Scrollbar(txt_previa, command=txt_previa.yview)
scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
txt_previa.config(yscrollcommand=scrollbar.set)

img_label = tk.Label(janela)
img_label.place(relx=0.78, rely=0.3, anchor=tk.CENTER)

# Botão para selecionar algum arquvio
btn_selecionar_arquivo = tk.Button(
    janela,
    text="Selecionar Arquivo",
    command=selecionar_arquivo,
    width=20,
    height=2,
    bg="#FFFFFF",
    fg="#000000",
    relief=tk.RAISED,
    bd=5,
)
btn_selecionar_arquivo.pack(padx=10, pady=10, anchor="w")


# botão para selecionar alguma imagem
btn_selecionar_imagem = tk.Button(
    janela,
    text="Selecionar Imagem",
    command=selecionar_imagem,
    width=20,
    height=2,
    bg="#FFFFFF",
    fg="#000000",
    relief=tk.RAISED,
    bd=5,
)
btn_selecionar_imagem.pack(padx=10, pady=10, anchor="w")

# Botão para capturar imagem da webcam
btn_capturar_imagem = tk.Button(
    janela,
    text="Camera",
    command=abrir_webcam,
    width=20,
    height=2,
    bg="#FFFFFF",
    fg="#000000",
    relief=tk.RAISED,
    bd=5,
)
btn_capturar_imagem.pack(padx=10, pady=10, anchor="w")

#botão para tirar a foto
btn_tirar_foto = tk.Button(
    janela,
    text="Tirar foto",
    command= capturar_imagem,
    width=20,
    height=2,
    bg="#FFFFFF",
    fg="#000000",
    state=tk.DISABLED,
    relief=tk.RAISED,
    bd=5,
)
btn_tirar_foto.pack(padx=10, pady=10, anchor="w")

# botão aplicar restrito apenas quando for selecionado algum tipo de arquvio/ imagem
btn_aplicar = tk.Button(
    janela,
    text="Aplicar",
    command=aplicar_selecao,
    width=20,
    height=2,
    bg="#FFFFFF",
    fg="#000000",
    state=tk.DISABLED,
    relief=tk.RAISED,
    bd=5,
)
btn_aplicar.pack(padx=10, pady=10, anchor="w")

# botão de créditos
btn_créditos = tk.Button(
    janela,
    text="Créditos",
    command=creditos_finais,
    width=20, # Aumente a largura para esticar o botão
    height=1, # Reduza a altura para torná-lo menor
    bg="#FFFFFF",
    fg="#000000",
    relief=tk.RAISED,
    bd=5,
)
btn_créditos.place(x=500, y=580, anchor=tk.CENTER)


#Colocando o texto na tela principal 
coopyright_label = tk.Label(janela, text="Copyright - Vinicius Cosmos  © 2023", font=("Arial", 10), fg="white", bg="#000000")
coopyright_label.place(x=10, y=altura_janela - 20)

janela.mainloop()